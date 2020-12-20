from docx import Document
from docx.shared import Inches
import pyttsx3

pyttsx3.speak('Hello')

document = Document()

# profile picture
document.add_picture(
    'picture.jpg',
    width=Inches(2.0)
)

# name phone number and email details
name = input('What is your name? ')
phone_number = input('What is your phone number? ')
email = input('What is your e-mail? ')

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email)

# about me
document.add_heading('About me')
document.add_paragraph(
    input('Tell me about yourself? ')
    )

# work experience
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter company ')
from_date = input('From Date ')
to_date = input('To Date ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input(
    'Describe your experience at ' + company)
p.add_run(experience_details)

# more experiences
while True:
    has_more_experiences = input(
        'Do you have more experiences? Yes or No ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter company ')
        from_date = input('From Date ')
        to_date = input('To Date ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input(
            'Describe your experience at ' + company + ' ')
        p.add_run(experience_details)
    else: 
        break

# skills
document.add_heading('Skills')
skill = input('Enter skill')
p = document.add_paragraph(skill)
p.style = 'List bullet'

while True:
    has_more_skill = input('Do you have more skill? Yes or No ')
    if has_more_skill.lower() == 'yes':
        skill = input('Enter skill')
        p = document.add_paragraph(skill)
        p.style = 'List bullet'
    else:
        break

# footer
section = document.section[0]
footer = document.footer
p = footer.paragraphs[0]
p.text = "CV generated using Amigoscode"

document.save('cv.docx')